import streamlit as st
from audio_recorder_streamlit import audio_recorder
import speech_recognition as sr
from docx import Document
from docx.shared import Pt
import base64
from io import BytesIO
import tempfile
import time

st.set_page_config(page_title="🎙️ Yes Sir- Incremental Stenographer", layout="centered")
st.title("🗣️ Voice Typing — Hindi / English / Hinglish")

# Initialize doc in session
if "doc" not in st.session_state:
    st.session_state.doc = Document()

# Select language
lang_option = st.selectbox("Select Language:", ["English", "Hindi", "Hinglish"])
lang_code = {"English": "en-IN", "Hindi": "hi-IN", "Hinglish": "hi-IN"}[lang_option]

# Choose input method
mode = st.radio("Input Method", ["🎤 Record Microphone", "📁 Upload WAV File"])

text = ""

# AUDIO RECORDING
if mode == "🎤 Record Microphone":
    st.markdown("### Step 2: Start Speaking")
    st.info("⏱️ Max recording time: ~60 seconds. Please finish within the limit.")

    audio_bytes = audio_recorder(pause_threshold=60.0)

    if audio_bytes:
        st.audio(audio_bytes, format="audio/wav")
        with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp:
            tmp.write(audio_bytes)
            tmp_path = tmp.name

        r = sr.Recognizer()
        with sr.AudioFile(tmp_path) as source:
            audio_data = r.record(source)

        try:
            text = r.recognize_google(audio_data, language=lang_code)
            st.success("✅ Transcription complete")
        except sr.UnknownValueError:
            st.error("❌ Could not understand audio")
            text = ""

# UPLOAD AUDIO FILE
if mode == "📁 Upload WAV File":
    uploaded = st.file_uploader("Upload .wav file", type=["wav"])
    if uploaded is not None:
        st.audio(uploaded, format="audio/wav")

        with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp:
            tmp.write(uploaded.read())
            tmp_path = tmp.name

        r = sr.Recognizer()
        with sr.AudioFile(tmp_path) as source:
            audio_data = r.record(source)

        try:
            text = r.recognize_google(audio_data, language=lang_code)
            st.success("✅ Transcription complete")
        except sr.UnknownValueError:
            st.error("❌ Could not understand audio")
            text = ""

# If transcription available
if text:
    edited = st.text_area("📝 Edit transcription", text, height=200)

    # Save to session's growing Word doc
    if st.button("➕ Add to session document"):
        run = st.session_state.doc.add_paragraph().add_run(edited)
        run.font.size = Pt(14)
        if lang_option != "English":
            run.font.name = "Mangal"
        st.success("✅ Added to session document!")

    # Allow download
    if st.button("💾 Download Full Document (.docx)"):
        buf = BytesIO()
        st.session_state.doc.save(buf)
        buf.seek(0)
        b64 = base64.b64encode(buf.read()).decode()
        href = f'<a href="data:application/octet-stream;base64,{b64}" download="session_transcript.doc">📥 Download .doc</a>'
        st.markdown(href, unsafe_allow_html=True)
